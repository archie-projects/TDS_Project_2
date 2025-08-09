import os
import json
import base64
import tempfile
import io
import re
from typing import Any, Dict, List, Optional, Union
import asyncio
import aiofiles
import httpx
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from bs4 import BeautifulSoup
import duckdb
from PIL import Image
import requests
import PyPDF2
from fastapi import FastAPI, File, UploadFile, HTTPException, Form, Query
from fastapi.responses import JSONResponse, HTMLResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
import uvicorn
import logging
from pathlib import Path
import traceback
from datetime import datetime
import google.generativeai as genai
from scipy import stats
from io import StringIO
import mimetypes

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Configuration
class Config:
    GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    REQUEST_TIMEOUT = 300  # 5 minutes

    @classmethod
    def validate(cls):
        if not cls.GEMINI_API_KEY:
            logger.critical("GEMINI_API_KEY environment variable not set. The application will not work.")
            raise ValueError("GEMINI_API_KEY not set")

# Initialize Gemini
if Config.GEMINI_API_KEY:
    genai.configure(api_key=Config.GEMINI_API_KEY)

class LLMProvider:
    """Abstract base class for LLM providers"""
    async def generate_response(self, prompt: str, json_mode: bool = False) -> str:
        raise NotImplementedError

class GeminiProvider(LLMProvider):
    def __init__(self, api_key: str):
        self.api_key = api_key
        # Use the gemini-pro model as requested
        self.model = genai.GenerativeModel('gemini-1.5-flash')

    async def generate_response(self, prompt: str, json_mode: bool = False) -> str:
        try:
            if json_mode:
                prompt += "\n\nIMPORTANT: Respond with valid JSON only, no other text or markdown."

            # For gemini-pro, json_mode is handled via instruction, not a direct API parameter
            response = self.model.generate_content(prompt)
            return response.text
        except Exception as e:
            logger.error(f"Gemini API error: {e}")
            raise

class WebSearchProvider:
    """Handles web search functionality"""
    
    @staticmethod
    async def search_web(query: str, num_results: int = 5) -> List[Dict[str, Any]]:
        """Search the web for information"""
        try:
            # Use multiple search strategies
            results = []
            
            # Strategy 1: Direct web scraping for specific domains
            if any(domain in query.lower() for domain in ['wikipedia', 'github', 'stackoverflow']):
                direct_results = await WebSearchProvider._direct_search(query)
                results.extend(direct_results)
            
            # Strategy 2: General web search using search engines
            search_results = await WebSearchProvider._general_search(query, num_results)
            results.extend(search_results)
            
            return results[:num_results]
            
        except Exception as e:
            logger.error(f"Web search failed: {e}")
            return [{"error": f"Search failed: {str(e)}", "query": query}]
    
    @staticmethod
    async def _direct_search(query: str) -> List[Dict[str, Any]]:
        """Direct search for specific domains"""
        results = []
        try:
            # Extract URLs from query if present
            urls = re.findall(r'https?://[^\s\)]+', query)
            
            for url in urls:
                scraped = await WebScraper.scrape_url(url)
                if 'error' not in scraped:
                    results.append({
                        "title": scraped.get('title', 'Unknown Title'),
                        "url": url,
                        "content": scraped.get('text', '')[:1000],
                        "type": "direct_scrape"
                    })
        except Exception as e:
            logger.error(f"Direct search error: {e}")
        
        return results
    
    @staticmethod
    async def _general_search(query: str, num_results: int) -> List[Dict[str, Any]]:
        """General web search using search APIs or scraping"""
        results = []
        try:
            # Use DuckDuckGo search (no API key required)
            search_url = f"https://html.duckduckgo.com/html/?q={query.replace(' ', '+')}"
            
            async with httpx.AsyncClient(timeout=30.0) as client:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                response = await client.get(search_url, headers=headers)
                
                if response.status_code == 200:
                    soup = BeautifulSoup(response.text, 'html.parser')
                    
                    # Extract search results
                    result_links = soup.find_all('a', class_='result__a')[:num_results]
                    
                    for link in result_links:
                        title = link.get_text().strip()
                        url = link.get('href', '')
                        
                        if url.startswith('/l/?uddg='):
                            # DuckDuckGo redirect URL, extract actual URL
                            import urllib.parse
                            url = urllib.parse.unquote(url.split('uddg=')[1])
                        
                        # Get snippet from result
                        result_div = link.find_parent('div', class_='result')
                        snippet = ""
                        if result_div:
                            snippet_elem = result_div.find('div', class_='result__snippet')
                            if snippet_elem:
                                snippet = snippet_elem.get_text().strip()
                        
                        results.append({
                            "title": title,
                            "url": url,
                            "content": snippet,
                            "type": "search_result"
                        })
        
        except Exception as e:
            logger.error(f"General search error: {e}")
        
        return results

class FileProcessor:
    """Handles different file types and extracts data"""

    @staticmethod
    async def process_file(file_path: str, filename: str) -> Dict[str, Any]:
        """Process uploaded file and extract data"""
        try:
            file_info = {
                "filename": filename,
                "type": None,
                "data": None,
                "metadata": {}
            }

            # Get file extension and MIME type
            file_ext = Path(filename).suffix.lower()
            mime_type, _ = mimetypes.guess_type(filename)

            # Handle various file types
            if file_ext in ['.csv', '.tsv'] or (mime_type and 'csv' in mime_type):
                file_info["type"] = "tabular"
                # Try different encodings and separators
                encodings = ['utf-8', 'latin-1', 'cp1252']
                separators = [',', '\t', ';', '|']
                
                for encoding in encodings:
                    for sep in separators:
                        try:
                            df = pd.read_csv(file_path, encoding=encoding, sep=sep, on_bad_lines='skip')
                            if df.shape[1] > 1:  # Valid table
                                file_info["data"] = df
                                file_info["metadata"] = {
                                    "rows": len(df),
                                    "columns": list(df.columns),
                                    "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
                                    "encoding": encoding,
                                    "separator": sep
                                }
                                break
                        except Exception:
                            continue
                    if file_info["data"] is not None:
                        break

            elif file_ext in ['.xlsx', '.xls'] or (mime_type and 'spreadsheet' in mime_type):
                file_info["type"] = "tabular"
                try:
                    df = pd.read_excel(file_path, engine='openpyxl' if file_ext == '.xlsx' else 'xlrd')
                    file_info["data"] = df
                    file_info["metadata"] = {
                        "rows": len(df),
                        "columns": list(df.columns),
                        "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()}
                    }
                except Exception:
                    # Try reading all sheets
                    try:
                        sheets = pd.read_excel(file_path, sheet_name=None)
                        # Use the first sheet with data
                        for sheet_name, df in sheets.items():
                            if not df.empty:
                                file_info["data"] = df
                                file_info["metadata"] = {
                                    "rows": len(df),
                                    "columns": list(df.columns),
                                    "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
                                    "sheet_name": sheet_name
                                }
                                break
                    except Exception as e:
                        logger.error(f"Excel processing error: {e}")

            elif file_ext == '.json' or (mime_type and 'json' in mime_type):
                file_info["type"] = "json"
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    file_info["data"] = data
                    file_info["metadata"] = {"keys": list(data.keys()) if isinstance(data, dict) else "array"}
                except Exception:
                    # Try with different encodings
                    for encoding in ['latin-1', 'cp1252']:
                        try:
                            with open(file_path, 'r', encoding=encoding) as f:
                                data = json.load(f)
                            file_info["data"] = data
                            file_info["metadata"] = {"keys": list(data.keys()) if isinstance(data, dict) else "array"}
                            break
                        except Exception:
                            continue

            elif file_ext == '.pdf' or (mime_type and 'pdf' in mime_type):
                file_info["type"] = "text"
                text = await FileProcessor._extract_pdf_text(file_path)
                file_info["data"] = text
                file_info["metadata"] = {"length": len(text), "pages": text.count('\n\n') + 1}

            elif file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp'] or (mime_type and 'image' in mime_type):
                file_info["type"] = "image"
                try:
                    with open(file_path, 'rb') as f:
                        img_data = f.read()
                    file_info["data"] = base64.b64encode(img_data).decode()
                    img = Image.open(file_path)
                    file_info["metadata"] = {"size": img.size, "format": img.format, "mode": img.mode}
                except Exception as e:
                    logger.error(f"Image processing error: {e}")

            elif file_ext in ['.txt', '.md', '.rst'] or (mime_type and 'text' in mime_type):
                # Try to read as a table first, then as plain text
                encodings = ['utf-8', 'latin-1', 'cp1252']
                
                for encoding in encodings:
                    try:
                        # First attempt: structured data
                        df = pd.read_csv(file_path, encoding=encoding, sep=None, engine='python', on_bad_lines='skip')
                        if df.shape[1] > 1:
                            file_info["type"] = "tabular"
                            file_info["data"] = df
                            file_info["metadata"] = {
                                "rows": len(df),
                                "columns": list(df.columns),
                                "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
                                "encoding": encoding
                            }
                            break
                    except Exception:
                        pass
                
                # If not tabular, read as plain text
                if file_info["data"] is None:
                    for encoding in encodings:
                        try:
                            with open(file_path, 'r', encoding=encoding) as f:
                                text = f.read()
                            file_info["type"] = "text"
                            file_info["data"] = text
                            file_info["metadata"] = {"length": len(text), "encoding": encoding}
                            break
                        except Exception:
                            continue

            elif file_ext in ['.xml', '.html', '.htm']:
                file_info["type"] = "markup"
                encodings = ['utf-8', 'latin-1', 'cp1252']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            content = f.read()
                        
                        # Parse with BeautifulSoup
                        soup = BeautifulSoup(content, 'html.parser' if file_ext in ['.html', '.htm'] else 'xml')
                        text = soup.get_text()
                        
                        file_info["data"] = text
                        file_info["metadata"] = {
                            "length": len(text),
                            "encoding": encoding,
                            "tags": len(soup.find_all()) if soup else 0
                        }
                        break
                    except Exception:
                        continue

            elif file_ext in ['.zip', '.rar', '.7z']:
                file_info["type"] = "archive"
                file_info["metadata"] = {"note": "Archive files require extraction before processing"}

            elif file_ext in ['.doc', '.docx'] or (mime_type and 'word' in mime_type):
                file_info["type"] = "document"
                try:
                    if file_ext == '.docx':
                        from docx import Document
                        doc = Document(file_path)
                        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                        file_info["data"] = text
                        file_info["metadata"] = {"length": len(text), "paragraphs": len(doc.paragraphs)}
                    else:
                        file_info["metadata"] = {"note": "DOC files require additional processing"}
                except ImportError:
                    file_info["metadata"] = {"note": "python-docx library required for DOCX processing"}
                except Exception as e:
                    logger.error(f"Document processing error: {e}")

            else:
                # Unknown file type - try to process as binary or text
                file_info["type"] = "unknown"
                try:
                    # Try to read as text first
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read(1000)  # Read first 1000 chars
                    if text.isprintable():
                        file_info["type"] = "text"
                        file_info["data"] = text
                        file_info["metadata"] = {"note": "Unknown text file type", "preview_length": len(text)}
                except Exception:
                    # Binary file
                    with open(file_path, 'rb') as f:
                        data = f.read(100)  # Read first 100 bytes
                    file_info["metadata"] = {"note": "Binary file - not processed", "size_bytes": os.path.getsize(file_path)}

            return file_info

        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            return {"filename": filename, "type": "error", "error": str(e)}

    @staticmethod
    async def _extract_pdf_text(file_path: str) -> str:
        """Extract text from PDF"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""

class WebScraper:
    """Handles web scraping with various strategies"""

    @staticmethod
    async def scrape_url(url: str) -> Dict[str, Any]:
        """Scrape data from URL"""
        try:
            # Clean URL first - remove any trailing characters
            url = url.rstrip(')').rstrip('.')

            async with httpx.AsyncClient(timeout=60.0) as client:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                response = await client.get(url, follow_redirects=True, headers=headers)
                response.raise_for_status()

            soup = BeautifulSoup(response.text, 'html.parser')

            # Try to find tables
            tables = soup.find_all('table')
            scraped_data = {
                "url": url,
                "title": soup.title.string if soup.title else "",
                "tables": [],
                "text": soup.get_text()[:5000],  # First 5000 chars
                "raw_html": response.text[:10000]  # First 10k chars
            }

            # Convert tables to dataframes
            for i, table in enumerate(tables[:5]):  # Limit to 5 tables
                try:
                    df = pd.read_html(StringIO(str(table)))[0]
                    scraped_data["tables"].append({
                        "index": i,
                        "columns": list(df.columns),
                        "rows": len(df),
                        "data": df.to_dict('records')[:100]  # Limit rows
                    })
                except Exception as e:
                    logger.warning(f"Could not parse table {i}: {e}")

            return scraped_data

        except Exception as e:
            logger.error(f"Error scraping URL {url}: {e}")
            return {"url": url, "error": str(e)}

class DataAnalystAgent:
    """Main agent that orchestrates the analysis"""

    def __init__(self, llm_provider: LLMProvider):
        self.llm = llm_provider
        self.temp_dir = tempfile.mkdtemp()

    def __del__(self):
        """Cleanup temp directory"""
        try:
            import shutil
            if hasattr(self, 'temp_dir') and os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
                logger.info(f"Cleaned up temp directory: {self.temp_dir}")
        except Exception as e:
            logger.warning(f"Failed to cleanup temp directory: {e}")

    async def analyze(self, questions: str, files: List[UploadFile] = None) -> Any:
        """Main analysis pipeline - now handles cases with no files"""
        try:
            # Step 1: Process all files (if any)
            processed_files = []
            if files:
                processed_files = await self._process_files(files)

            # Step 2: Extract URLs from questions and scrape if needed
            urls = re.findall(r'https?://[^\s\)]+', questions)
            scraped_data = []
            for url in urls:
                scraped = await WebScraper.scrape_url(url)
                scraped_data.append(scraped)

            # Step 3: Perform web search if no files and no URLs
            search_results = []
            if not processed_files and not scraped_data:
                # Determine if this needs a web search
                if self._needs_web_search(questions):
                    search_query = self._extract_search_query(questions)
                    search_results = await WebSearchProvider.search_web(search_query)
                    logger.info(f"Performed web search for: {search_query}")

            # Step 4: Create analysis context
            context = {
                "files": processed_files,
                "scraped_data": scraped_data,
                "search_results": search_results,
                "questions": questions,
                "has_data": bool(processed_files or scraped_data or search_results)
            }

            # Step 5: Generate execution plan
            plan = await self._create_execution_plan(context)

            # Step 6: Execute the plan
            execution_results = await self._execute_plan(plan, context)

            # Step 7: Generate final response
            final_answer = await self._generate_final_answer(questions, execution_results, context, plan)

            return final_answer

        except Exception as e:
            logger.error(f"Analysis failed: {e}")
            logger.error(traceback.format_exc())
            raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

    def _needs_web_search(self, questions: str) -> bool:
        """Determine if the question needs web search"""
        # Keywords that indicate current information needs
        search_indicators = [
            'current', 'latest', 'recent', 'today', 'news', 'weather',
            'stock price', 'market', 'covid', 'election', 'trending',
            'what is', 'who is', 'when did', 'where is', 'how to'
        ]
        
        questions_lower = questions.lower()
        return any(indicator in questions_lower for indicator in search_indicators)

    def _extract_search_query(self, questions: str) -> str:
        """Extract search query from questions"""
        # Clean up the questions to create a good search query
        query = questions.strip()
        
        # Remove question marks and common question words
        query = re.sub(r'\?', '', query)
        query = re.sub(r'^(what|how|when|where|who|why|which)\s+', '', query, flags=re.IGNORECASE)
        
        # Take first sentence if multiple sentences
        query = query.split('.')[0].split('\n')[0]
        
        return query.strip()

    async def _process_files(self, files: List[UploadFile]) -> List[Dict]:
        """Process all uploaded files"""
        processed = []

        for file in files:
            if file.filename == "questions.txt":
                continue

            # Save file temporarily
            file_path = os.path.join(self.temp_dir, file.filename)
            async with aiofiles.open(file_path, 'wb') as f:
                content = await file.read()
                await f.write(content)

            # Process file
            file_info = await FileProcessor.process_file(file_path, file.filename)
            processed.append(file_info)

        return processed

    async def _create_execution_plan(self, context: Dict) -> Dict:
        """Create a flexible execution plan based on context"""

        data_summary = ""
        if context['files']:
            data_summary += f"Files: {len(context['files'])} files\n"
            for f in context['files']:
                data_summary += f"  - {f.get('filename', 'unknown')} ({f.get('type', 'unknown')})\n"
        
        if context['scraped_data']:
            data_summary += f"Scraped URLs: {len(context['scraped_data'])}\n"
        
        if context['search_results']:
            data_summary += f"Web Search Results: {len(context['search_results'])}\n"

        plan_prompt = f"""You are an expert AI assistant. Analyze the user's questions and available data to create a comprehensive execution plan.

QUESTIONS TO ANSWER:
{context['questions']}

AVAILABLE DATA:
{data_summary}

DATA HAS BEEN PROVIDED: {context['has_data']}

Create a JSON execution plan to answer the questions. The primary data table for queries will be named 'main_data'.

1. If NO data is provided and questions need factual answers:
   - Use "knowledge_response" or "web_search" actions.

2. If data IS provided (files, URLs, or search results):
   - Create steps for data loading, analysis, and visualization.
   - For any calculations, data extraction, or quantitative questions, use the "run_sql" action. You should write a SQL query to be executed against the 'main_data' table.
   - For statistical analysis (like correlation or regression), use "statistical_analysis".
   - For plotting data, use "create_visualization".

Response format:
{{
    "analysis_type": "knowledge_based|data_analysis|web_research|mixed",
    "expected_response_format": "text|json_object|mixed",
    "response_structure": "describe the expected structure",
    "steps": [
        {{
            "step_id": "step_1",
            "action": "knowledge_response|web_search|load_data|run_sql|statistical_analysis|create_visualization|text_analysis",
            "description": "clear description of what this step does",
            "params": {{
                "query": "SQL query, search query, or question",
                "data_source": "specific file or search results",
                "analysis_type": "descriptive|statistical|comparative",
                "output_format": "text|json|dataframe|visualization"
            }},
            "expected_output": "description of expected output"
        }}
    ]
}}

Respond with valid JSON only."""

        try:
            plan_text = await self.llm.generate_response(plan_prompt, json_mode=True)
            # Clean up response if needed
            plan_text = re.sub(r'^```json\s*', '', plan_text)
            plan_text = re.sub(r'\s*```$', '', plan_text)

            plan = json.loads(plan_text)
            logger.info(f"Generated execution plan: {json.dumps(plan, indent=2)}")
            return plan
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse plan JSON: {e}")
            logger.error(f"Raw response: {plan_text}")
            # Fallback plan
            return {
                "analysis_type": "knowledge_based" if not context['has_data'] else "data_analysis",
                "expected_response_format": "text",
                "response_structure": "Direct answer to the questions",
                "steps": [
                    {
                        "step_id": "step_1", 
                        "action": "knowledge_response" if not context['has_data'] else "load_data", 
                        "description": "Provide answer based on available information", 
                        "params": {"query": context['questions']}
                    }
                ]
            }

    def _clean_dataframe(self, df: pd.DataFrame) -> pd.DataFrame:
        """Clean and standardize dataframe"""
        # Clean column names
        df.columns = [re.sub(r'[^A-Za-z0-9_]', '', str(col).strip().replace(' ', '_').replace('(', '').replace(')', '')) for col in df.columns]

        for col in df.columns:
            # Attempt to convert to datetime for date-like columns
            if df[col].dtype == 'object':
                if any(term in col.lower() for term in ['date', 'time', 'year']):
                    try:
                        # Coerce errors will turn unparseable values into NaT
                        converted_dates = pd.to_datetime(df[col], errors='coerce')
                        # Only convert if a significant portion of the column is date-like
                        if converted_dates.notna().sum() / len(df[col]) > 0.5:
                            df[col] = converted_dates
                            continue
                    except Exception:
                        pass # Ignore if conversion fails

            # Clean and convert numeric columns
            # Check if column is not purely descriptive text
            if df[col].dtype == 'object':
                try:
                    # Remove currency symbols, commas, and percentage signs for conversion
                    cleaned_series = df[col].astype(str).str.replace(r'[$,%\s]', '', regex=True)
                    numeric_series = pd.to_numeric(cleaned_series, errors='coerce')
                    
                    # Only convert if the column becomes mostly numeric
                    if numeric_series.notna().sum() / len(df[col]) > 0.7:
                        df[col] = numeric_series
                except Exception:
                    pass # Ignore if cleaning/conversion fails

        return df

    def _clean_sql_query(self, query: str, available_columns: List[str]) -> str:
        """Clean and validate SQL query"""
        # Replace common column name patterns with actual column names
        for col in available_columns:
            # Case-insensitive replacement
            pattern = re.compile(re.escape(col), re.IGNORECASE)
            query = pattern.sub(f'"{col}"', query)

        # Ensure we're querying the main table
        if 'FROM' not in query.upper() and 'from' not in query.lower():
            query = query + " FROM main_data"

        return query

    async def _perform_statistical_analysis(self, df: pd.DataFrame, params: Dict) -> Dict:
        """Perform statistical analysis"""
        try:
            analysis_type = params.get('analysis_type', 'correlation')
            columns = params.get('columns', [])
            if not columns: # Try to get from 'variables' if 'columns' is empty
                columns = params.get('variables', [])


            # Convert datetime columns to numeric for regression if needed
            numeric_df = df.copy()
            for col in numeric_df.columns:
                if pd.api.types.is_datetime64_any_dtype(numeric_df[col]):
                    numeric_df[col] = numeric_df[col].apply(lambda x: x.toordinal() if pd.notna(x) else np.nan)

            if analysis_type == 'regression' and len(columns) >= 2:
                x_col, y_col = columns[0], columns[1]

                if x_col in numeric_df.columns and y_col in numeric_df.columns:
                    # Remove NaN values
                    clean_data = numeric_df[[x_col, y_col]].dropna()
                    if len(clean_data) > 1:
                        slope, intercept, r_value, p_value, std_err = stats.linregress(
                            clean_data[x_col], clean_data[y_col]
                        )
                        return {
                            'type': 'regression',
                            'x_column': x_col,
                            'y_column': y_col,
                            'slope': float(slope),
                            'intercept': float(intercept),
                            'r_squared': float(r_value**2),
                            'p_value': float(p_value),
                            'sample_size': len(clean_data)
                        }

            elif analysis_type == 'correlation':
                # If specific columns are provided for correlation, use them
                if len(columns) >= 2:
                    numeric_cols = [col for col in columns if col in numeric_df.select_dtypes(include=np.number).columns]
                else: # Otherwise, use all numeric columns
                    numeric_cols = numeric_df.select_dtypes(include=[np.number]).columns.tolist()

                if len(numeric_cols) >= 2:
                    corr_matrix = numeric_df[numeric_cols].corr()

                    # Find strongest correlation (excluding diagonal)
                    corr_values = corr_matrix.abs().unstack()
                    corr_values = corr_values[corr_values < 1.0]  # Remove diagonal
                    if not corr_values.empty:
                        strongest_idx = corr_values.idxmax()
                        strongest_value = corr_values.max()
                    else:
                        strongest_idx = ('N/A', 'N/A')
                        strongest_value = 0

                    return {
                        'type': 'correlation',
                        'correlation_matrix': corr_matrix.to_dict(),
                        'strongest_correlation': {
                            'columns': strongest_idx,
                            'value': float(strongest_value)
                        },
                        'numeric_columns': numeric_cols
                    }

            return {
                'type': 'statistical_analysis',
                'error': f'Cannot perform {analysis_type} analysis with provided params.',
                'available_columns': list(df.columns)
            }

        except Exception as e:
            logger.error(f"Statistical analysis error: {e}")
            return {'type': 'statistical_analysis', 'error': str(e)}

    async def _create_visualization(self, df: pd.DataFrame, params: Dict) -> Dict:
        """Create visualizations"""
        try:
            plot_type = params.get('plot_type', 'bar')
            # Look in multiple common keys for column names
            x_col = params.get('x_col') or params.get('x_axis')
            y_col = params.get('y_col') or params.get('y_axis')
            
            # If still not found, check the generic 'columns' key
            if not x_col and params.get('columns') and len(params['columns']) > 0:
                x_col = params['columns'][0]
            if not y_col and params.get('columns') and len(params['columns']) > 1:
                y_col = params['columns'][1]

            aggregation = params.get('aggregation', 'sum')

            # Auto-detect columns if they are not explicitly provided
            if not x_col:
                # Default to the first categorical column for x-axis
                cat_cols = df.select_dtypes(include=['object', 'category']).columns
                if len(cat_cols) > 0:
                    x_col = cat_cols[0]
                else: # Or the first column overall if no categorical
                    x_col = df.columns[0]
            
            if not y_col and plot_type != 'bar': # Bar can be a value_counts plot
                # Default to the first numeric column for y-axis
                num_cols = df.select_dtypes(include=np.number).columns
                if len(num_cols) > 0:
                    y_col = num_cols[0]


            if not x_col or x_col not in df.columns:
                return {'type': 'visualization', 'error': f'X-axis column "{x_col}" not found. Available columns: {list(df.columns)}'}

            # Create plot
            plt.figure(figsize=(12, 8))
            plt.style.use('default')

            plot_data = df.copy()
            # Drop NaNs from plotting columns to avoid errors
            plot_cols = [c for c in [x_col, y_col] if c is not None]
            plot_data.dropna(subset=plot_cols, inplace=True)
            if plot_data.empty:
                return {'type': 'visualization', 'error': 'No data available for plotting after removing missing values.'}


            # Create the actual plot
            if plot_type == 'bar':
                if y_col and y_col in plot_data.columns:
                     # Group by x_col and aggregate y_col
                    if x_col in plot_data.columns:
                        agg_func = {'sum': 'sum', 'count': 'count', 'mean': 'mean'}.get(aggregation, 'sum')
                        plot_values = plot_data.groupby(x_col)[y_col].agg(agg_func).sort_values(ascending=False)
                        sns.barplot(x=plot_values.index, y=plot_values.values, palette="viridis")
                    else:
                        return {'type': 'visualization', 'error': f'X-axis column "{x_col}" not found for bar plot aggregation.'}
                else:
                    # If no y_col, do a count plot
                    sns.countplot(x=x_col, data=plot_data, palette="viridis", order = plot_data[x_col].value_counts().index)

            elif plot_type == 'scatter':
                 if not y_col or y_col not in df.columns:
                     return {'type': 'visualization', 'error': f'Y-axis column "{y_col}" is required for a scatter plot.'}
                 sns.scatterplot(x=x_col, y=y_col, data=plot_data, alpha=0.7)
                 if params.get('add_regression', False) or params.get('regression_line', False):
                     sns.regplot(x=x_col, y=y_col, data=plot_data, scatter=False, color='red')


            elif plot_type == 'line':
                if not y_col or y_col not in df.columns:
                     return {'type': 'visualization', 'error': f'Y-axis column "{y_col}" is required for a line plot.'}
                # If x-axis is datetime, seaborn handles it nicely
                sns.lineplot(x=x_col, y=y_col, data=plot_data, marker='o')

            # Styling
            plt.xlabel(str(x_col).replace('_', ' ').title(), fontsize=14)
            plt.ylabel(str(y_col if y_col else 'Count').replace('_', ' ').title(), fontsize=14)
            plt.xticks(rotation=45, ha='right')

            # Create meaningful title
            title_y = y_col if y_col else 'Count'
            title = f"{plot_type.title()} of {str(title_y).replace('_',' ').title()} by {str(x_col).replace('_',' ').title()}"
            plt.title(title, fontsize=16, fontweight='bold', pad=20)

            plt.grid(True, alpha=0.3, linestyle='--')
            plt.tight_layout()


            # Save to base64
            buffer = io.BytesIO()
            plt.savefig(buffer, format='png', dpi=100) # Use a reasonable DPI
            plt.close()
            img_data = base64.b64encode(buffer.getvalue()).decode()


            return {
                'type': 'visualization',
                'format': 'base64_png',
                'data': f"data:image/png;base64,{img_data}",
                'size_bytes': len(img_data),
                'plot_info': {
                    'type': plot_type,
                    'x_column': x_col,
                    'y_column': y_col,
                    'title': title
                }
            }

        except Exception as e:
            logger.error(f"Visualization error: {e}")
            logger.error(traceback.format_exc())
            return {'type': 'visualization', 'error': str(e)}

    async def _execute_plan(self, plan: Dict, context: Dict) -> Dict:
        """Execute the analysis plan"""
        results = {}

        # Initialize DuckDB connection
        con = duckdb.connect(database=':memory:', read_only=False)
        try:
            main_df = None

            # First, handle local files if any exist
            for file_info in context['files']:
                if file_info.get('type') == 'tabular' and file_info.get('data') is not None:
                    try:
                        df = file_info['data']
                        df = self._clean_dataframe(df)
                        table_name = f"file_{re.sub(r'[^A-Za-z0-9_]', '', file_info['filename'])}"
                        con.register(table_name, df)
                        if main_df is None:
                            main_df = df
                            con.register('main_data', df)
                            logger.info(f"Registered file '{file_info['filename']}' as main_data with shape: {df.shape}")
                            logger.info(f"Columns: {list(df.columns)}")
                    except Exception as e:
                        logger.error(f"Error processing file {file_info['filename']}: {e}")

            # Handle scraped data
            for scraped in context['scraped_data']:
                if 'tables' in scraped and scraped['tables']:
                    for table_info in scraped['tables']:
                        try:
                            # Use the first valid table as the main dataframe
                            if main_df is None and table_info.get('data'):
                                df = pd.DataFrame(table_info['data'])
                                df = self._clean_dataframe(df)
                                main_df = df
                                con.register('main_data', main_df)
                                logger.info(f"Registered scraped table as main_data with shape: {df.shape}")
                                logger.info(f"Columns: {list(df.columns)}")
                                break # Use the first table and stop
                        except Exception as e:
                            logger.error(f"Error processing scraped table: {e}")

            # Execute plan steps
            for step in plan.get('steps', []):
                step_id = step['step_id']
                action = step['action']
                params = step.get('params', {})

                logger.info(f"Executing {step_id}: {action} with params: {params}")

                try:
                    if action == 'knowledge_response':
                        query = params.get('query', context['questions'])
                        knowledge_prompt = f"Answer the following question based on your training knowledge: {query}"
                        answer = await self.llm.generate_response(knowledge_prompt)
                        results[step_id] = {
                            'type': 'knowledge_response',
                            'answer': answer
                        }

                    elif action == 'web_search':
                        query = params.get('query', context['questions'])
                        search_results = await WebSearchProvider.search_web(query)
                        search_content = "\n\n".join([f"Source: {r.get('title')}\nURL: {r.get('url')}\nContent: {r.get('content')}" for r in search_results if 'error' not in r])
                        synthesis_prompt = f"Based on these search results, provide a comprehensive answer to: {query}\n\nResults:\n{search_content}"
                        synthesized_answer = await self.llm.generate_response(synthesis_prompt)
                        results[step_id] = {
                            'type': 'web_search_response',
                            'answer': synthesized_answer,
                            'source': 'Web search'
                        }

                    elif action == 'load_data':
                        # This action is now implicitly handled by loading files/scraped data.
                        # We can add more advanced logic here if needed (e.g., from databases)
                        if main_df is not None:
                             results[step_id] = {'type': 'success', 'message': f'Data loaded as main_data, shape: {main_df.shape}, columns: {list(main_df.columns)}'}
                        else:
                             results[step_id] = {'type': 'error', 'error': 'No data could be loaded.'}


                    elif action == 'run_sql':
                        query = params.get('query', '')
                        if query and main_df is not None:
                            # The planner is instructed to use 'main_data' as the table name
                            logger.info(f"Executing SQL: {query}")
                            result_df = con.execute(query).df()
                            results[step_id] = {
                                'type': 'dataframe',
                                'data': result_df.to_dict('records'),
                                'columns': list(result_df.columns),
                                'shape': result_df.shape
                            }
                            logger.info(f"SQL result shape: {result_df.shape}")
                        else:
                            results[step_id] = {'type': 'error', 'error': 'No data available or no query provided'}

                    elif action == 'statistical_analysis':
                        if main_df is None:
                            results[step_id] = {'type': 'error', 'error': 'Statistical analysis requires a loaded dataframe.'}
                            continue
                        analysis_result = await self._perform_statistical_analysis(main_df.copy(), params)
                        results[step_id] = analysis_result

                    elif action == 'create_visualization':
                        if main_df is None:
                            results[step_id] = {'type': 'error', 'error': 'Visualization requires a loaded dataframe.'}
                            continue
                        plot_result = await self._create_visualization(main_df.copy(), params)
                        results[step_id] = plot_result

                    elif action == 'text_analysis':
                        text_content = ""
                        for file_info in context['files']:
                            if file_info.get('type') == 'text' and file_info.get('data'):
                                text_content += file_info['data'] + "\n\n"
                        for result in context.get('search_results', []):
                            if 'content' in result:
                                text_content += result['content'] + "\n\n"
                        
                        if text_content:
                            analysis_prompt = f"""Analyze the following text content and provide insights based on the user's question:

Question: {context['questions']}

Text Content (first 5000 chars):
{text_content[:5000]}

Provide a structured analysis addressing the user's questions."""

                            analysis = await self.llm.generate_response(analysis_prompt)
                            results[step_id] = {
                                'type': 'text_analysis',
                                'analysis': analysis
                            }
                        else:
                            results[step_id] = {'type': 'error', 'error': 'No text content available for analysis'}
                    
                    # This action is deprecated and removed as it's not generic.
                    # The 'run_sql' action should be used instead.
                    elif action == 'data_cleaning':
                        # Data is cleaned upon loading. This step is for acknowledgment.
                         if main_df is not None:
                            results[step_id] = {'type': 'success', 'message': 'Data cleaning applied during initial load.'}
                         else:
                            results[step_id] = {'type': 'error', 'error': 'No data to clean.'}

                except Exception as e:
                    logger.error(f"Error in step {step_id}: {action} - {e}")
                    logger.error(traceback.format_exc())
                    results[step_id] = {'type': 'error', 'error': str(e)}

        except Exception as e:
            logger.error(f"Error during plan execution: {e}")
            logger.error(traceback.format_exc())
            results['execution_error'] = {'type': 'error', 'error': str(e)}
        finally:
            con.close()

        return results

    async def _generate_final_answer(self, questions: str, results: Dict, context: Dict, plan: Dict) -> Any:
        """Generate the final answer in the requested format by synthesizing results."""

        # Handle simple cases first (no data analysis)
        if plan.get("analysis_type") in ["knowledge_based", "web_research"]:
            for step_id, result in results.items():
                if result.get('type') in ['knowledge_response', 'web_search_response']:
                    # Return the direct answer from these steps
                    return {"answer": result.get('answer', 'No answer found.'), "source": result.get('source', 'AI Assistant')}

        # For complex data analysis, synthesize all results.
        synthesis_prompt = f"""You are a data synthesis AI. Your task is to provide the final answer in the EXACT format requested by the user, based on the results of an execution plan.

ORIGINAL QUESTIONS:
{questions}

EXECUTION PLAN AND RESULTS:
{json.dumps(results, indent=2, default=str)}

CONTEXT:
- The user expected a response in this format: {plan.get('expected_response_format', 'Not specified')}
- With this structure: {plan.get('response_structure', 'Not specified')}

CRITICAL INSTRUCTIONS:
1. Carefully review the ORIGINAL QUESTIONS.
2. Examine the EXECUTION PLAN AND RESULTS. Each key is a step_id from the plan.
3. Synthesize the information from all successful steps to construct the final answer.
4. The final answer MUST be a single, valid JSON object that directly answers the user's questions.
5. If a step resulted in an error, state that the specific piece of information could not be retrieved and mention the error if possible.
6. If the results contain a `dataframe` type, extract the key values from the `data` list inside it.
7. If the results contain a `visualization` type, include the 'data' field (the base64 string) in the final JSON output under an appropriate key.

Generate the final JSON object now."""

        try:
            final_response_str = await self.llm.generate_response(synthesis_prompt, json_mode=True)
            # Clean up response (remove markdown code blocks)
            final_response_str = re.sub(r'^```json\s*', '', final_response_str)
            final_response_str = re.sub(r'\s*```$', '', final_response_str)

            # Try to parse as JSON to validate
            parsed_json = json.loads(final_response_str)
            return parsed_json

        except Exception as e:
            logger.error(f"Final answer synthesis failed: {e}")
            logger.error(f"LLM Raw Output for Synthesis:\n{final_response_str}")

            # Fallback: Return a structured error and all partial results
            return {
                "error": "Failed to generate a synthesized JSON response.",
                "detail": "The AI's final output could not be parsed as valid JSON.",
                "llm_output": final_response_str,
                "execution_results": results # Provide all successful results for debugging
            }

# Initialize FastAPI app
app = FastAPI(
    title="Universal Data Analyst Agent",
    description="AI-powered data analysis with support for multiple file formats, web scraping, and knowledge-based responses",
    version="1.1.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Validate configuration on startup
try:
    Config.validate()
except ValueError as e:
    logger.critical(f"Configuration error: {e}")
    # Exit if essential config is missing
    exit(1)

# Initialize LLM provider
def get_llm_provider():
    """Initializes and returns the Gemini LLM provider."""
    return GeminiProvider(Config.GEMINI_API_KEY)

# API Routes
@app.get("/", response_class=HTMLResponse)
async def root():
    """Root endpoint that serves the main HTML interface"""
    if Path("index.html").exists():
        return FileResponse("index.html")
    else:
        return HTMLResponse("""
        <html>
            <body>
                <h1>Universal Data Analyst Agent</h1>
                <p>AI-powered data analysis with support for:</p>
                <ul>
                    <li>Multiple file formats (CSV, Excel, PDF, JSON, Images, etc.)</li>
                    <li>Web scraping and URL analysis</li>
                    <li>Knowledge-based question answering</li>
                    <li>Statistical analysis and visualizations</li>
                </ul>
                <p>API Documentation: <a href="/docs">/docs</a></p>
                <p>Health Check: <a href="/health">/health</a></p>
            </body>
        </html>
        """)

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "llm_provider": "gemini",
        "version": "1.1.0",
        "capabilities": [
            "file_processing",
            "web_scraping", 
            "knowledge_qa",
            "data_analysis",
            "visualization",
            "sql_execution"
        ]
    }

@app.get("/config")
async def get_config():
    """Get system configuration"""
    return {
        "llm_provider": "gemini",
        "llm_model": "gemini-1.5-flash",
        "max_file_size_mb": Config.MAX_FILE_SIZE // (1024 * 1024),
        "request_timeout_seconds": Config.REQUEST_TIMEOUT,
        "has_gemini_key": bool(Config.GEMINI_API_KEY),
        "supported_file_types": [
            "csv", "tsv", "xlsx", "xls", "json", "pdf", "txt", "md",
            "png", "jpg", "jpeg", "gif", "bmp", "tiff", "webp",
            "xml", "html", "htm", "docx", "zip", "rar", "7z"
        ]
    }

@app.post("/api/")
async def analyze_data(files: List[UploadFile] = File(default=[]), questions: str = Form(default="")):
    """Main analysis endpoint - now handles requests without files"""
    start_time = datetime.now()

    try:
        # Extract questions from different sources
        questions_content = questions
        data_files = []

        # Check if questions come from uploaded files
        for file in files:
            if file.filename == "questions.txt":
                questions_content = (await file.read()).decode('utf-8')
            else:
                # Reset file pointer for other files
                await file.seek(0)
                data_files.append(file)

        # If no questions provided via form or file, return error
        if not questions_content.strip():
            raise HTTPException(status_code=400, detail="No questions provided. Use 'questions' form field or upload 'questions.txt' file.")

        # Validate file sizes
        for file in data_files:
            # file.size is an optional int, so we check if it exists
            if hasattr(file, 'size') and file.size and file.size > Config.MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=413,
                    detail=f"File {file.filename} exceeds maximum size of {Config.MAX_FILE_SIZE // (1024*1024)}MB"
                )

        # Initialize agent and analyze
        llm_provider = get_llm_provider()
        agent = DataAnalystAgent(llm_provider)

        logger.info(f"Starting analysis with {len(data_files)} files")
        logger.info(f"Questions: {questions_content[:200]}...")

        result = await agent.analyze(questions_content, data_files if data_files else None)

        processing_time = (datetime.now() - start_time).total_seconds()
        logger.info(f"Analysis completed in {processing_time:.2f} seconds")

        return JSONResponse(content=result)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Analysis failed: {e}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

@app.post("/api/ask")
async def ask_question(question: str = Form(...)):
    """Simple question endpoint without file uploads"""
    return await analyze_data(files=[], questions=question)

@app.post("/api/search")
async def web_search(query: str = Form(...), num_results: int = Form(default=5)):
    """Web search endpoint"""
    try:
        results = await WebSearchProvider.search_web(query, num_results)
        return JSONResponse(content={"query": query, "results": results})
    except Exception as e:
        logger.error(f"Search failed: {e}")
        raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")

@app.post("/upload-test")
async def upload_test(files: List[UploadFile] = File(...)):
    """Test endpoint to check file uploads and processing"""
    file_info = []
    for file in files:
        content = await file.read()
        
        # Try to determine file type
        file_ext = Path(file.filename).suffix.lower()
        file_type = "unknown"
        
        if file_ext in ['.csv', '.tsv']:
            file_type = "tabular"
        elif file_ext in ['.xlsx', '.xls']:
            file_type = "spreadsheet"
        elif file_ext == '.json':
            file_type = "json"
        elif file_ext == '.pdf':
            file_type = "pdf"
        elif file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
            file_type = "image"
        elif file_ext in ['.txt', '.md']:
            file_type = "text"
        
        file_info.append({
            "filename": file.filename,
            "content_type": file.content_type,
            "size": len(content),
            "detected_type": file_type,
            "extension": file_ext
        })
    
    return {"uploaded_files": file_info}

# Error handlers
@app.exception_handler(413)
async def file_too_large_handler(request, exc):
    return JSONResponse(
        status_code=413,
        content={"error": "File too large", "max_size_mb": Config.MAX_FILE_SIZE // (1024*1024)}
    )

@app.exception_handler(500)
async def internal_server_error_handler(request, exc):
    logger.error(f"Internal server error: {exc}")
    return JSONResponse(
        status_code=500,
        content={"error": "Internal server error", "detail": str(exc)}
    )

# Development server runner
if __name__ == "__main__":
    import sys

    # Check for required dependencies
    try:
        import uvicorn
    except ImportError:
        print("uvicorn not found. Install with: pip install uvicorn")
        sys.exit(1)

    # Set up logging for development
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )

    print("\n🚀 Starting Universal Data Analyst Agent...")
    print(f"📊 LLM Provider: Gemini")
    print(f"🤖 AI Model: gemini-1.5-flash")
    print(f"📁 Max File Size: {Config.MAX_FILE_SIZE // (1024*1024)}MB")
    print(f"⏱️  Request Timeout: {Config.REQUEST_TIMEOUT}s")
    print("\n🌐 Server will be available at: http://localhost:8000")
    print("📖 API Documentation: http://localhost:8000/docs")
    print("🔍 Health Check: http://localhost:8000/health")
    print("\n✨ Key Improvements in this Version:")
    print("   - Generic analysis using AI-generated SQL queries.")
    print("   - Removed hardcoded logic for more flexible data handling.")
    print("   - Improved final response synthesis and error handling.")
    print("   - Enhanced visualization and statistical analysis functions.")

    uvicorn.run(
        "app:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        access_log=True
    )
