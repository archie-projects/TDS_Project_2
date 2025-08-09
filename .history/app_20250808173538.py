import os
import json
import base64
import tempfile
import io
import re
from typing import Any, Dict, List, Optional, Union
import asyncio
import aiofiles
import httpx
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from bs4 import BeautifulSoup
import duckdb
from PIL import Image
import requests
import PyPDF2
from fastapi import FastAPI, File, UploadFile, HTTPException, Form, Query
from fastapi.responses import JSONResponse, HTMLResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
import uvicorn
import logging
from pathlib import Path
import traceback
from datetime import datetime
import google.generativeai as genai
from scipy import stats
from io import StringIO
import mimetypes

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- Configuration Class (No Changes) ---
class Config:
    GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    REQUEST_TIMEOUT = 300  # 5 minutes

    @classmethod
    def validate(cls):
        if not cls.GEMINI_API_KEY:
            logger.critical("GEMINI_API_KEY environment variable not set. The application will not work.")
            raise ValueError("GEMINI_API_KEY not set")

# --- Gemini Initialization (No Changes) ---
if Config.GEMINI_API_KEY:
    genai.configure(api_key=Config.GEMINI_API_KEY)

# --- LLMProvider and GeminiProvider Classes (No Changes) ---
class LLMProvider:
    async def generate_response(self, prompt: str, json_mode: bool = False) -> str:
        raise NotImplementedError

class GeminiProvider(LLMProvider):
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.model = genai.GenerativeModel('gemini-1.5-flash')

    async def generate_response(self, prompt: str, json_mode: bool = False) -> str:
        try:
            generation_config = {}
            if json_mode:
                generation_config={"response_mime_type": "application/json"}
            
            # The gemini-1.5-flash model supports setting the response format directly.
            response = self.model.generate_content(prompt, generation_config=generation_config)
            return response.text
        except Exception as e:
            logger.error(f"Gemini API error: {e}")
            raise

# --- WebSearchProvider, FileProcessor, WebScraper Classes (No Changes) ---
# These classes were functioning correctly for the described problem.
class WebSearchProvider:
    """Handles web search functionality"""
    
    @staticmethod
    async def search_web(query: str, num_results: int = 5) -> List[Dict[str, Any]]:
        """Search the web for information"""
        try:
            # Use multiple search strategies
            results = []
            
            # Strategy 1: Direct web scraping for specific domains
            if any(domain in query.lower() for domain in ['wikipedia', 'github', 'stackoverflow']):
                direct_results = await WebSearchProvider._direct_search(query)
                results.extend(direct_results)
            
            # Strategy 2: General web search using search engines
            search_results = await WebSearchProvider._general_search(query, num_results)
            results.extend(search_results)
            
            return results[:num_results]
            
        except Exception as e:
            logger.error(f"Web search failed: {e}")
            return [{"error": f"Search failed: {str(e)}", "query": query}]
    
    @staticmethod
    async def _direct_search(query: str) -> List[Dict[str, Any]]:
        """Direct search for specific domains"""
        results = []
        try:
            # Extract URLs from query if present
            urls = re.findall(r'https?://[^\s\)]+', query)
            
            for url in urls:
                scraped = await WebScraper.scrape_url(url)
                if 'error' not in scraped:
                    results.append({
                        "title": scraped.get('title', 'Unknown Title'),
                        "url": url,
                        "content": scraped.get('text', '')[:1000],
                        "type": "direct_scrape"
                    })
        except Exception as e:
            logger.error(f"Direct search error: {e}")
        
        return results
    
    @staticmethod
    async def _general_search(query: str, num_results: int) -> List[Dict[str, Any]]:
        """General web search using search APIs or scraping"""
        results = []
        try:
            # Use DuckDuckGo search (no API key required)
            search_url = f"https://html.duckduckgo.com/html/?q={query.replace(' ', '+')}"
            
            async with httpx.AsyncClient(timeout=30.0) as client:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                response = await client.get(search_url, headers=headers)
                
                if response.status_code == 200:
                    soup = BeautifulSoup(response.text, 'html.parser')
                    
                    # Extract search results
                    result_links = soup.find_all('a', class_='result__a')[:num_results]
                    
                    for link in result_links:
                        title = link.get_text().strip()
                        url = link.get('href', '')
                        
                        if url.startswith('/l/?uddg='):
                            # DuckDuckGo redirect URL, extract actual URL
                            import urllib.parse
                            url = urllib.parse.unquote(url.split('uddg=')[1])
                        
                        # Get snippet from result
                        result_div = link.find_parent('div', class_='result')
                        snippet = ""
                        if result_div:
                            snippet_elem = result_div.find('div', class_='result__snippet')
                            if snippet_elem:
                                snippet = snippet_elem.get_text().strip()
                        
                        results.append({
                            "title": title,
                            "url": url,
                            "content": snippet,
                            "type": "search_result"
                        })
        
        except Exception as e:
            logger.error(f"General search error: {e}")
        
        return results

class FileProcessor:
    """Handles different file types and extracts data"""

    @staticmethod
    async def process_file(file_path: str, filename: str) -> Dict[str, Any]:
        """Process uploaded file and extract data"""
        try:
            file_info = {
                "filename": filename,
                "type": None,
                "data": None,
                "metadata": {}
            }

            # Get file extension and MIME type
            file_ext = Path(filename).suffix.lower()
            mime_type, _ = mimetypes.guess_type(filename)

            # Handle various file types
            if file_ext in ['.csv', '.tsv'] or (mime_type and 'csv' in mime_type):
                file_info["type"] = "tabular"
                # Try different encodings and separators
                encodings = ['utf-8', 'latin-1', 'cp1252']
                separators = [',', '\t', ';', '|']
                
                for encoding in encodings:
                    for sep in separators:
                        try:
                            df = pd.read_csv(file_path, encoding=encoding, sep=sep, on_bad_lines='skip')
                            if df.shape[1] > 1:  # Valid table
                                file_info["data"] = df
                                file_info["metadata"] = {
                                    "rows": len(df),
                                    "columns": list(df.columns),
                                    "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
                                    "encoding": encoding,
                                    "separator": sep
                                }
                                break
                        except Exception:
                            continue
                    if file_info["data"] is not None:
                        break

            elif file_ext in ['.xlsx', '.xls'] or (mime_type and 'spreadsheet' in mime_type):
                file_info["type"] = "tabular"
                try:
                    df = pd.read_excel(file_path, engine='openpyxl' if file_ext == '.xlsx' else 'xlrd')
                    file_info["data"] = df
                    file_info["metadata"] = {
                        "rows": len(df),
                        "columns": list(df.columns),
                        "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()}
                    }
                except Exception:
                    # Try reading all sheets
                    try:
                        sheets = pd.read_excel(file_path, sheet_name=None)
                        # Use the first sheet with data
                        for sheet_name, df in sheets.items():
                            if not df.empty:
                                file_info["data"] = df
                                file_info["metadata"] = {
                                    "rows": len(df),
                                    "columns": list(df.columns),
                                    "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
                                    "sheet_name": sheet_name
                                }
                                break
                    except Exception as e:
                        logger.error(f"Excel processing error: {e}")

            elif file_ext == '.json' or (mime_type and 'json' in mime_type):
                file_info["type"] = "json"
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    file_info["data"] = data
                    file_info["metadata"] = {"keys": list(data.keys()) if isinstance(data, dict) else "array"}
                except Exception:
                    # Try with different encodings
                    for encoding in ['latin-1', 'cp1252']:
                        try:
                            with open(file_path, 'r', encoding=encoding) as f:
                                data = json.load(f)
                            file_info["data"] = data
                            file_info["metadata"] = {"keys": list(data.keys()) if isinstance(data, dict) else "array"}
                            break
                        except Exception:
                            continue

            elif file_ext == '.pdf' or (mime_type and 'pdf' in mime_type):
                file_info["type"] = "text"
                text = await FileProcessor._extract_pdf_text(file_path)
                file_info["data"] = text
                file_info["metadata"] = {"length": len(text), "pages": text.count('\n\n') + 1}

            elif file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp'] or (mime_type and 'image' in mime_type):
                file_info["type"] = "image"
                try:
                    with open(file_path, 'rb') as f:
                        img_data = f.read()
                    file_info["data"] = base64.b64encode(img_data).decode()
                    img = Image.open(file_path)
                    file_info["metadata"] = {"size": img.size, "format": img.format, "mode": img.mode}
                except Exception as e:
                    logger.error(f"Image processing error: {e}")

            elif file_ext in ['.txt', '.md', '.rst'] or (mime_type and 'text' in mime_type):
                # Try to read as a table first, then as plain text
                encodings = ['utf-8', 'latin-1', 'cp1252']
                
                for encoding in encodings:
                    try:
                        # First attempt: structured data
                        df = pd.read_csv(file_path, encoding=encoding, sep=None, engine='python', on_bad_lines='skip')
                        if df.shape[1] > 1:
                            file_info["type"] = "tabular"
                            file_info["data"] = df
                            file_info["metadata"] = {
                                "rows": len(df),
                                "columns": list(df.columns),
                                "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
                                "encoding": encoding
                            }
                            break
                    except Exception:
                        pass
                
                # If not tabular, read as plain text
                if file_info["data"] is None:
                    for encoding in encodings:
                        try:
                            with open(file_path, 'r', encoding=encoding) as f:
                                text = f.read()
                            file_info["type"] = "text"
                            file_info["data"] = text
                            file_info["metadata"] = {"length": len(text), "encoding": encoding}
                            break
                        except Exception:
                            continue

            elif file_ext in ['.xml', '.html', '.htm']:
                file_info["type"] = "markup"
                encodings = ['utf-8', 'latin-1', 'cp1252']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            content = f.read()
                        
                        # Parse with BeautifulSoup
                        soup = BeautifulSoup(content, 'html.parser' if file_ext in ['.html', '.htm'] else 'xml')
                        text = soup.get_text()
                        
                        file_info["data"] = text
                        file_info["metadata"] = {
                            "length": len(text),
                            "encoding": encoding,
                            "tags": len(soup.find_all()) if soup else 0
                        }
                        break
                    except Exception:
                        continue

            elif file_ext in ['.zip', '.rar', '.7z']:
                file_info["type"] = "archive"
                file_info["metadata"] = {"note": "Archive files require extraction before processing"}

            elif file_ext in ['.doc', '.docx'] or (mime_type and 'word' in mime_type):
                file_info["type"] = "document"
                try:
                    if file_ext == '.docx':
                        from docx import Document
                        doc = Document(file_path)
                        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                        file_info["data"] = text
                        file_info["metadata"] = {"length": len(text), "paragraphs": len(doc.paragraphs)}
                    else:
                        file_info["metadata"] = {"note": "DOC files require additional processing"}
                except ImportError:
                    file_info["metadata"] = {"note": "python-docx library required for DOCX processing"}
                except Exception as e:
                    logger.error(f"Document processing error: {e}")

            else:
                # Unknown file type - try to process as binary or text
                file_info["type"] = "unknown"
                try:
                    # Try to read as text first
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read(1000)  # Read first 1000 chars
                    if text.isprintable():
                        file_info["type"] = "text"
                        file_info["data"] = text
                        file_info["metadata"] = {"note": "Unknown text file type", "preview_length": len(text)}
                except Exception:
                    # Binary file
                    with open(file_path, 'rb') as f:
                        data = f.read(100)  # Read first 100 bytes
                    file_info["metadata"] = {"note": "Binary file - not processed", "size_bytes": os.path.getsize(file_path)}

            return file_info

        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            return {"filename": filename, "type": "error", "error": str(e)}

    @staticmethod
    async def _extract_pdf_text(file_path: str) -> str:
        """Extract text from PDF"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""

class WebScraper:
    """Handles web scraping with various strategies"""

    @staticmethod
    async def scrape_url(url: str) -> Dict[str, Any]:
        """Scrape data from URL"""
        try:
            # Clean URL first - remove any trailing characters
            url = url.rstrip(')').rstrip('.')

            async with httpx.AsyncClient(timeout=60.0) as client:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                response = await client.get(url, follow_redirects=True, headers=headers)
                response.raise_for_status()

            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Try to find all tables, which is common for data extraction
            tables = soup.find_all('table', {'class': 'wikitable'})
            if not tables:
                tables = soup.find_all('table')

            scraped_data = {
                "url": url,
                "title": soup.title.string if soup.title else "",
                "tables": [],
                "text": soup.get_text()[:5000],  # First 5000 chars
                "raw_html": response.text[:10000]  # First 10k chars
            }

            # Convert tables to dataframes
            for i, table in enumerate(tables[:5]):  # Limit to 5 tables
                try:
                    # pd.read_html is great for parsing HTML tables
                    df_list = pd.read_html(StringIO(str(table)), flavor='bs4')
                    if df_list:
                        df = df_list[0]
                         # Clean up MultiIndex columns if they exist
                        if isinstance(df.columns, pd.MultiIndex):
                            df.columns = ['_'.join(map(str, col)).strip() for col in df.columns.values]
                        
                        scraped_data["tables"].append({
                            "index": i,
                            "columns": list(df.columns),
                            "rows": len(df),
                            "data": df.to_dict('records')[:100]  # Limit rows
                        })
                except Exception as e:
                    logger.warning(f"Could not parse table {i} from {url}: {e}")

            return scraped_data

        except Exception as e:
            logger.error(f"Error scraping URL {url}: {e}")
            return {"url": url, "error": str(e)}

# --- DataAnalystAgent Class (MAJOR CHANGES) ---
class DataAnalystAgent:
    """Main agent that orchestrates the analysis"""

    def __init__(self, llm_provider: LLMProvider):
        self.llm = llm_provider
        self.temp_dir = tempfile.mkdtemp()

    def __del__(self):
        """Cleanup temp directory"""
        try:
            import shutil
            if hasattr(self, 'temp_dir') and os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
                logger.info(f"Cleaned up temp directory: {self.temp_dir}")
        except Exception as e:
            logger.warning(f"Failed to cleanup temp directory: {e}")

    async def analyze(self, questions: str, files: List[UploadFile] = None) -> Any:
        """Main analysis pipeline"""
        try:
            processed_files = []
            if files:
                processed_files = await self._process_files(files)

            urls = re.findall(r'https?://[^\s\)]+', questions)
            scraped_data = []
            for url in urls:
                scraped = await WebScraper.scrape_url(url)
                scraped_data.append(scraped)

            search_results = []
            if not processed_files and not scraped_data and self._needs_web_search(questions):
                search_query = self._extract_search_query(questions)
                search_results = await WebSearchProvider.search_web(search_query)
                logger.info(f"Performed web search for: {search_query}")

            context = {
                "files": processed_files,
                "scraped_data": scraped_data,
                "search_results": search_results,
                "questions": questions,
                "has_data": bool(processed_files or scraped_data or search_results)
            }

            plan = await self._create_execution_plan(context)
            execution_results = await self._execute_plan(plan, context)
            final_answer = await self._generate_final_answer(questions, execution_results, plan)

            return final_answer

        except Exception as e:
            logger.error(f"Analysis failed: {e}")
            logger.error(traceback.format_exc())
            raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

    def _needs_web_search(self, questions: str) -> bool:
        """Determine if the question needs web search"""
        search_indicators = [
            'current', 'latest', 'recent', 'today', 'news', 'weather',
            'stock price', 'market', 'covid', 'election', 'trending',
            'what is', 'who is', 'when did', 'where is', 'how to'
        ]
        questions_lower = questions.lower()
        return any(indicator in questions_lower for indicator in search_indicators) and not re.search(r'https?://', questions_lower)


    def _extract_search_query(self, questions: str) -> str:
        query = questions.strip()
        query = re.sub(r'\?', '', query)
        query = re.sub(r'^(what|how|when|where|who|why|which)\s+', '', query, flags=re.IGNORECASE)
        query = query.split('.')[0].split('\n')[0]
        return query.strip()

    async def _process_files(self, files: List[UploadFile]) -> List[Dict]:
        processed = []
        for file in files:
            if file.filename == "questions.txt":
                continue
            file_path = os.path.join(self.temp_dir, file.filename)
            async with aiofiles.open(file_path, 'wb') as f:
                content = await file.read()
                await f.write(content)
            file_info = await FileProcessor.process_file(file_path, file.filename)
            processed.append(file_info)
        return processed

    async def _create_execution_plan(self, context: Dict) -> Dict:
        """[FIXED] Create a robust execution plan using only valid actions."""
        
        data_summary = ""
        if context['files']:
            for f in context['files']:
                data_summary += f"- File: {f.get('filename', 'unknown')} (Type: {f.get('type', 'unknown')}, Columns: {f.get('metadata', {}).get('columns')})\n"
        if context['scraped_data']:
            for s in context['scraped_data']:
                 for t in s.get('tables', []):
                    data_summary += f"- Scraped Table from {s.get('url')}: (Columns: {t.get('columns')})\n"
        if context['search_results']:
            data_summary += f"- Web Search Results: {len(context['search_results'])} results\n"
        if not data_summary:
            data_summary = "No data provided."

        # This prompt is heavily constrained to prevent hallucination.
        plan_prompt = f"""
You are an expert data analysis planner. Your job is to create a JSON execution plan to answer the user's questions based on the available data.

**CRITICAL: You can ONLY use the following actions in your plan:**
- `run_sql`: To query, filter, aggregate, and transform tabular data. Use this for all calculations.
- `statistical_analysis`: For statistical tests like correlation or regression.
- `create_visualization`: To generate plots.
- `knowledge_response`: To answer questions that don't require data analysis (e.g., "hello").

**User Questions:**
{context['questions']}

**Available Data Summary:**
{data_summary}

**Instructions:**
1.  Carefully analyze each question.
2.  For each question, create one or more steps in the plan.
3.  Use `run_sql` for any data manipulation, filtering, or calculation. The data is in a table called `main_data`. You can use DuckDB SQL functions for cleaning (e.g., `regexp_replace`, `CAST`).
4.  For statistical questions (like correlation or regression slope), use `statistical_analysis`.
5.  For plotting questions, use `create_visualization`.
6.  Each step must have a unique `step_id` (e.g., `q1_1`, `q2_1`).

**Output Format (JSON only):**
{{
    "steps": [
        {{
            "step_id": "...",
            "action": "run_sql | statistical_analysis | create_visualization | knowledge_response",
            "description": "...",
            "params": {{ ... }}
        }}
    ]
}}

**Example Plan for "What is the total revenue?":**
{{
    "steps": [
        {{
            "step_id": "q1_total_revenue",
            "action": "run_sql",
            "description": "Calculate the sum of the revenue column.",
            "params": {{
                "query": "SELECT SUM(revenue) FROM main_data"
            }}
        }}
    ]
}}
"""
        try:
            plan_text = await self.llm.generate_response(plan_prompt, json_mode=True)
            plan = json.loads(plan_text)
            logger.info(f"Generated execution plan: {json.dumps(plan, indent=2)}")
            return plan
        except (json.JSONDecodeError, TypeError) as e:
            logger.error(f"Failed to parse plan JSON: {e}\nRaw response was: {plan_text}")
            # Create a fallback plan if JSON parsing fails
            return {"steps": [{"step_id": "fallback_knowledge", "action": "knowledge_response", "params": {"query": context['questions']}}]}

    def _clean_dataframe(self, df: pd.DataFrame) -> pd.DataFrame:
        """Clean and standardize dataframe columns."""
        # A more robust column cleaner
        def clean_col(col):
            # Remove characters that are bad for SQL
            new_col = re.sub(r'[^\w\s]', '', str(col))
            # Replace spaces with underscores
            new_col = re.sub(r'\s+', '_', new_col)
            # Remove leading/trailing underscores
            new_col = new_col.strip('_')
            return new_col

        df.columns = [clean_col(col) for col in df.columns]
        return df

    async def _perform_statistical_analysis(self, df: pd.DataFrame, params: Dict) -> Dict:
        """[FIXED] Perform statistical analysis based on parameters."""
        try:
            analysis_type = params.get('analysis_type', 'correlation')
            columns = params.get('columns', [])
            
            if not columns or len(columns) < 2:
                return {'error': f"Statistical analysis requires at least 2 columns. Got: {columns}"}
            
            # Ensure columns exist in the dataframe
            for col in columns:
                if col not in df.columns:
                    return {'error': f"Column '{col}' not found in dataframe for analysis. Available columns: {list(df.columns)}"}

            # Drop rows with missing values in the relevant columns
            clean_df = df[columns].dropna()
            
            if len(clean_df) < 2:
                return {'error': "Not enough data points for statistical analysis after dropping NaNs."}

            x = pd.to_numeric(clean_df[columns[0]], errors='coerce')
            y = pd.to_numeric(clean_df[columns[1]], errors='coerce')
            
            # Reclean after numeric conversion
            valid_data = pd.concat([x, y], axis=1).dropna()
            
            if len(valid_data) < 2:
                 return {'error': "Not enough valid numeric data points for analysis."}


            if analysis_type == 'correlation':
                correlation, p_value = stats.pearsonr(valid_data[columns[0]], valid_data[columns[1]])
                return {'type': 'correlation', 'correlation': float(correlation), 'p_value': float(p_value)}

            elif analysis_type == 'regression':
                slope, intercept, r_value, p_value, std_err = stats.linregress(
                    valid_data[columns[0]], valid_data[columns[1]]
                )
                return {
                    'type': 'regression', 'slope': float(slope), 'intercept': float(intercept),
                    'r_squared': float(r_value**2), 'p_value': float(p_value)
                }
                
            return {'error': f"Unknown analysis type: {analysis_type}"}

        except Exception as e:
            logger.error(f"Statistical analysis error: {e}")
            return {'type': 'statistical_analysis', 'error': str(e)}

    async def _create_visualization(self, df: pd.DataFrame, params: Dict) -> Dict:
        """[FIXED] Create visualizations with precise styling."""
        try:
            plot_type = params.get('plot_type', 'scatter')
            x_col = params.get('x_col')
            y_col = params.get('y_col')
            
            if not x_col or not y_col:
                return {'error': "Visualization requires 'x_col' and 'y_col' parameters."}
            
            if x_col not in df.columns or y_col not in df.columns:
                return {'error': f"Columns for plotting not found. Need {x_col}, {y_col}. Available: {list(df.columns)}"}

            plot_df = df[[x_col, y_col]].copy()
            # Convert to numeric, coercing errors to NaN
            plot_df[x_col] = pd.to_numeric(plot_df[x_col], errors='coerce')
            plot_df[y_col] = pd.to_numeric(plot_df[y_col], errors='coerce')
            plot_df.dropna(inplace=True)

            if plot_df.empty:
                return {'error': "No valid data to plot after cleaning."}

            plt.figure(figsize=(10, 6)) # Standard figure size
            
            if plot_type == 'scatter':
                plt.scatter(plot_df[x_col], plot_df[y_col], alpha=0.6)

                if params.get('add_regression'):
                    # Match prompt: 'dotted red regression line'
                    line_color = params.get('regression_color', 'red')
                    line_style = params.get('regression_style', 'dotted')
                    
                    slope, intercept, _, _, _ = stats.linregress(plot_df[x_col], plot_df[y_col])
                    line_x = np.array(plt.xlim())
                    line_y = slope * line_x + intercept
                    plt.plot(line_x, line_y, color=line_color, linestyle=line_style, linewidth=2, label='Regression Line')
                    plt.legend()
            
            plt.xlabel(x_col.replace('_', ' ').title())
            plt.ylabel(y_col.replace('_', ' ').title())
            plt.title(params.get('title', f'{y_col.title()} vs. {x_col.title()}'))
            plt.grid(True, alpha=0.3)
            plt.tight_layout()

            buffer = io.BytesIO()
            # Save with settings to control size, aiming for under 100kB
            plt.savefig(buffer, format='png', dpi=90, bbox_inches='tight')
            plt.close()
            
            img_data = base64.b64encode(buffer.getvalue()).decode()
            img_uri = f"data:image/png;base64,{img_data}"
            
            # Final check on size
            if len(img_uri) > 100000:
                logger.warning(f"Image size {len(img_uri)} is over 100kB limit. May be rejected.")

            return {
                'type': 'visualization',
                'format': 'base64_png',
                'data': img_uri,
                'size_bytes': len(img_uri)
            }
        except Exception as e:
            logger.error(f"Visualization error: {e}\n{traceback.format_exc()}")
            return {'type': 'visualization', 'error': str(e)}

    async def _execute_plan(self, plan: Dict, context: Dict) -> Dict:
        """[FIXED] Execute a robust plan with error handling."""
        results = {}
        main_df = None
        con = duckdb.connect(database=':memory:', read_only=False)
        con.execute("INSTALL httpfs; LOAD httpfs;")
        con.execute("INSTALL parquet; LOAD parquet;")

        try:
            # Consolidate all tabular data into one main dataframe
            all_dfs = []
            for file_info in context['files']:
                if isinstance(file_info.get('data'), pd.DataFrame):
                    all_dfs.append(self._clean_dataframe(file_info['data']))
            
            for scraped in context['scraped_data']:
                for table in scraped.get('tables', []):
                    df = pd.DataFrame(table['data'])
                    all_dfs.append(self._clean_dataframe(df))
            
            if all_dfs:
                # For simplicity, we'll use the largest table as the main one.
                # A more complex agent might join or union them.
                main_df = max(all_dfs, key=len)
                con.register('main_data', main_df)
                logger.info(f"Registered 'main_data' with shape: {main_df.shape} and columns: {list(main_df.columns)}")

            # Execute plan steps
            for step in plan.get('steps', []):
                step_id = step['step_id']
                action = step['action']
                params = step.get('params', {})
                logger.info(f"Executing {step_id}: {action}")

                try:
                    if action == 'knowledge_response':
                        prompt = f"Answer the following question: {params.get('query', context['questions'])}"
                        results[step_id] = {'type': 'knowledge', 'answer': await self.llm.generate_response(prompt)}
                    
                    elif main_df is None and action in ['run_sql', 'statistical_analysis', 'create_visualization']:
                         results[step_id] = {'error': f"Action '{action}' requires data, but none was loaded."}
                         continue

                    elif action == 'run_sql':
                        query = params.get('query')
                        if not query:
                             results[step_id] = {'error': "No query provided for run_sql action."}
                             continue
                        logger.info(f"Running SQL: {query}")
                        result = con.execute(query).df()
                        results[step_id] = {'type': 'sql_result', 'data': result.to_dict('records')}
                    
                    elif action == 'statistical_analysis':
                        results[step_id] = await self._perform_statistical_analysis(main_df, params)
                    
                    elif action == 'create_visualization':
                        results[step_id] = await self._create_visualization(main_df, params)
                    
                    else:
                        results[step_id] = {'error': f"Unknown or unsupported action: {action}"}

                except Exception as e:
                    logger.error(f"Error executing step {step_id}: {e}\n{traceback.format_exc()}")
                    results[step_id] = {'error': str(e)}
        finally:
            con.close()
        
        logger.info(f"Execution results: {json.dumps(results, indent=2, default=str)}")
        return results

    async def _generate_final_answer(self, questions: str, results: Dict, plan: Dict) -> Any:
        """[FIXED] Generate the final answer in the requested format."""

        # This prompt is much simpler. It just assembles the pre-computed results.
        synthesis_prompt = f"""
You are an AI assistant that formats final answers. Your only job is to assemble the provided results into the format requested by the user.

**User's Original Request:**
